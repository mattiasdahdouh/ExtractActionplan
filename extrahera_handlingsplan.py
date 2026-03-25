import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import csv
import threading
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


def extract_table_from_docx(filepath):
    """Extraherar den första tabellen (7 rader) från ett Word-dokument."""
    doc = Document(filepath)
    if not doc.tables:
        return None

    table = doc.tables[0]
    data = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip()
            value = cells[1].text.strip()
            if key:
                data[key] = value
    return data


def unique_path(directory, filename):
    """Returnerar en sökväg som inte redan finns — lägger till (1), (2) etc. vid konflikt."""
    base, ext = os.path.splitext(filename)
    candidate = os.path.join(directory, filename)
    counter = 1
    while os.path.exists(candidate):
        candidate = os.path.join(directory, f"{base} ({counter}){ext}")
        counter += 1
    return candidate


def save_to_excel(all_data, output_dir):
    """Sparar extraherad data till en Excel-fil."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Handlingsplaner"

    # Stilar
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1F497D")
    category_fill = PatternFill("solid", fgColor="DCE6F1")
    category_font = Font(bold=True, size=10)
    data_font = Font(size=10)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    if not all_data:
        return None

    # Samla alla unika kategorier (bevara ordning från första dokumentet)
    all_keys = list(all_data[0]["data"].keys()) if all_data else []

    # Skriv rubrikrad: Filnamn | Kategori1 | Kategori2 | ...
    ws.cell(row=1, column=1, value="Filnamn").font = header_font
    ws.cell(row=1, column=1).fill = header_fill
    ws.cell(row=1, column=1).alignment = center_align
    ws.cell(row=1, column=1).border = border

    for col_idx, key in enumerate(all_keys, start=2):
        cell = ws.cell(row=1, column=col_idx, value=key)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = border

    # Skriv data per fil
    for row_idx, entry in enumerate(all_data, start=2):
        # Filnamn
        cell = ws.cell(row=row_idx, column=1, value=entry["filename"])
        cell.font = category_font
        cell.fill = category_fill
        cell.alignment = left_align
        cell.border = border

        # Värden
        for col_idx, key in enumerate(all_keys, start=2):
            value = entry["data"].get(key, "")
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = data_font
            cell.alignment = left_align
            cell.border = border

    # Kolumnbredder
    ws.column_dimensions["A"].width = 30
    for col_idx in range(2, len(all_keys) + 2):
        ws.column_dimensions[get_column_letter(col_idx)].width = 22

    # Frys rubrikrad
    ws.freeze_panes = "A2"

    output_path = unique_path(output_dir, "handlingsplaner_sammanstallning.xlsx")
    wb.save(output_path)
    return output_path


def save_to_csv(all_data, output_dir):
    """Sparar extraherad data till en CSV-fil (UTF-8 med BOM för Excel-kompatibilitet)."""
    if not all_data:
        return None

    all_keys = list(all_data[0]["data"].keys())
    output_path = unique_path(output_dir, "handlingsplaner_sammanstallning.csv")

    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(["Filnamn"] + all_keys)
        for entry in all_data:
            row = [entry["filename"]] + [entry["data"].get(k, "") for k in all_keys]
            writer.writerow(row)

    return output_path


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Extrahera Handlingsplaner")
        self.resizable(False, False)
        self.configure(bg="#F0F0F0")

        self._build_ui()
        self._center_window(520, 340)

    def _center_window(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth() - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build_ui(self):
        pad = {"padx": 14, "pady": 6}

        title = tk.Label(
            self,
            text="Extrahera Handlingsplaner till Excel",
            font=("Segoe UI", 13, "bold"),
            bg="#F0F0F0",
            fg="#1F497D",
        )
        title.grid(row=0, column=0, columnspan=3, pady=(18, 10))

        # Källmapp
        tk.Label(self, text="Källmapp (Word-filer):", bg="#F0F0F0", font=("Segoe UI", 10)).grid(
            row=1, column=0, sticky="w", **pad
        )
        self.source_var = tk.StringVar()
        tk.Entry(self, textvariable=self.source_var, width=38, font=("Segoe UI", 9)).grid(
            row=1, column=1, sticky="ew", padx=(0, 4)
        )
        tk.Button(self, text="Bläddra", command=self._browse_source, width=9).grid(
            row=1, column=2, padx=(0, 14)
        )

        # Utmapp
        tk.Label(self, text="Utdatamapp (Excel):", bg="#F0F0F0", font=("Segoe UI", 10)).grid(
            row=2, column=0, sticky="w", **pad
        )
        self.output_var = tk.StringVar()
        tk.Entry(self, textvariable=self.output_var, width=38, font=("Segoe UI", 9)).grid(
            row=2, column=1, sticky="ew", padx=(0, 4)
        )
        tk.Button(self, text="Bläddra", command=self._browse_output, width=9).grid(
            row=2, column=2, padx=(0, 14)
        )

        # Statustext
        self.status_var = tk.StringVar(value="Välj mappar och klicka på Starta.")
        self.status_label = tk.Label(
            self,
            textvariable=self.status_var,
            bg="#F0F0F0",
            font=("Segoe UI", 9),
            fg="#444444",
            anchor="w",
        )
        self.status_label.grid(row=3, column=0, columnspan=3, sticky="w", padx=14, pady=(10, 2))

        # Progressbar
        self.progress = ttk.Progressbar(self, orient="horizontal", length=480, mode="determinate")
        self.progress.grid(row=4, column=0, columnspan=3, padx=14, pady=(0, 12))

        # Knappar
        btn_frame = tk.Frame(self, bg="#F0F0F0")
        btn_frame.grid(row=5, column=0, columnspan=3, pady=(4, 18))

        self.run_btn = tk.Button(
            btn_frame,
            text="Starta",
            command=self._start,
            width=14,
            bg="#1F497D",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            relief="flat",
            cursor="hand2",
        )
        self.run_btn.pack(side="left", padx=6)

        tk.Button(
            btn_frame,
            text="Avsluta",
            command=self.destroy,
            width=10,
            font=("Segoe UI", 10),
            relief="flat",
            cursor="hand2",
        ).pack(side="left", padx=6)

    def _browse_source(self):
        folder = filedialog.askdirectory(title="Välj mapp med Word-dokument")
        if folder:
            self.source_var.set(folder)

    def _browse_output(self):
        folder = filedialog.askdirectory(title="Välj mapp för Excel-fil")
        if folder:
            self.output_var.set(folder)

    def _start(self):
        source = self.source_var.get().strip()
        output = self.output_var.get().strip()

        if not source or not os.path.isdir(source):
            messagebox.showerror("Fel", "Välj en giltig källmapp.")
            return
        if not output or not os.path.isdir(output):
            messagebox.showerror("Fel", "Välj en giltig utdatamapp.")
            return

        self.run_btn.config(state="disabled")
        thread = threading.Thread(target=self._process, args=(source, output), daemon=True)
        thread.start()

    def _process(self, source, output):
        docx_files = [
            f for f in os.listdir(source)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]

        if not docx_files:
            self._set_status("Inga Word-dokument hittades i källmappen.", error=True)
            self.run_btn.config(state="normal")
            return

        total = len(docx_files)
        self.progress["maximum"] = total
        self.progress["value"] = 0

        all_data = []
        errors = []

        for i, filename in enumerate(docx_files, start=1):
            self._set_status(f"Bearbetar ({i}/{total}): {filename}")
            filepath = os.path.join(source, filename)
            try:
                data = extract_table_from_docx(filepath)
                if data:
                    all_data.append({"filename": filename, "data": data})
                else:
                    errors.append(f"{filename}: Ingen tabell hittad")
            except Exception as e:
                errors.append(f"{filename}: {e}")
            self.progress["value"] = i
            self.update_idletasks()

        if not all_data:
            self._set_status("Kunde inte extrahera data från några dokument.", error=True)
            self.run_btn.config(state="normal")
            return

        self._set_status("Sparar filer...")
        try:
            excel_path = save_to_excel(all_data, output)
            csv_path = save_to_csv(all_data, output)
        except Exception as e:
            self._set_status(f"Fel vid sparande: {e}", error=True)
            self.run_btn.config(state="normal")
            return

        msg = (
            f"Klar! {len(all_data)} dokument exporterade.\n\n"
            f"Excel: {excel_path}\nCSV:   {csv_path}"
        )
        if errors:
            msg += f"\n\nFel i {len(errors)} fil(er):\n" + "\n".join(errors)

        self._set_status(f"Klar! {len(all_data)} dokument exporterade (Excel + CSV).")
        self.run_btn.config(state="normal")
        messagebox.showinfo("Klar", msg)

    def _set_status(self, text, error=False):
        self.status_var.set(text)
        self.status_label.config(fg="#CC0000" if error else "#444444")
        self.update_idletasks()


if __name__ == "__main__":
    app = App()
    app.mainloop()
